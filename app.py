from flask import render_template, make_response, send_file
import pdfkit
from docx import Document
from bs4 import BeautifulSoup

from flask import Flask, jsonify
from main.routes import main
from contact.routes import contact
from cicd.routes import cicd
import platform
import os
import time

from models.models import db  # <-- Import SQLAlchemy db instance
from dotenv import load_dotenv
load_dotenv()

# ⬇️ NEW: Prometheus & psutil
from prometheus_flask_exporter import PrometheusMetrics
from prometheus_client import Gauge
import psutil

app = Flask(__name__)
app.secret_key = "devops-is-awesome"

# Setup Prometheus metrics
metrics = PrometheusMetrics(app)  # Adds default metrics at /metrics

# Custom resource usage metrics
memory_percent_gauge = Gauge('flask_memory_usage_percent', 'Memory usage percent')
cpu_percent_gauge = Gauge('flask_cpu_usage_percent', 'CPU usage percent')

# Update custom metrics on each request
@app.before_request
def update_resource_metrics():
    memory_percent_gauge.set(psutil.virtual_memory().percent)
    cpu_percent_gauge.set(psutil.cpu_percent(interval=None))

# SQLAlchemy Configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///cicd.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# Create database tables
with app.app_context():
    db.create_all()

# Register Blueprints
app.register_blueprint(main)
app.register_blueprint(contact)
app.register_blueprint(cicd)

# Track uptime
start_time = time.time()

# System Info Route
@app.route('/sysinfo')
def sysinfo():
    uptime = time.time() - start_time
    try:
        container_id = open("/etc/hostname").read().strip()
    except Exception:
        container_id = "Unavailable"

    memory = psutil.virtual_memory()
    cpu_percent = psutil.cpu_percent(interval=0.5)

    info = {
        "hostname": platform.node(),
        "system": platform.system(),
        "release": platform.release(),
        "python_version": platform.python_version(),
        "uptime_seconds": round(uptime),
        "container_id": container_id,
        "memory_total_MB": round(memory.total / (1024 ** 2)),
        "memory_used_MB": round(memory.used / (1024 ** 2)),
        "memory_percent": memory.percent,
        "cpu_percent": cpu_percent
    }
    return jsonify(info)
    
  
from flask import render_template, make_response
import pdfkit
from docx import Document
from bs4 import BeautifulSoup

@app.route("/download/pdf")
def generate_pdf():
    rendered = render_template("cv_template.html")
    # PDFKit config (you can set config= if wkhtmltopdf is not in PATH)
    pdf = pdfkit.from_string(rendered, False)
    response = make_response(pdf)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "attachment; filename=cv_pialroy.pdf"
    return response

@app.route("/download/docx")
def generate_docx():
    rendered = render_template("cv_template.html")
    soup = BeautifulSoup(rendered, "html.parser")
    
    doc = Document()
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li']):
        text = elem.get_text(strip=True)
        if elem.name.startswith('h'):
            doc.add_heading(text, level=int(elem.name[-1]))
        elif elem.name == 'p':
            doc.add_paragraph(text)
        elif elem.name == 'li':
            doc.add_paragraph(f'• {text}', style='List Bullet')

    response = make_response()
    doc.save("cv_pialroy.docx")
    with open("cv_pialroy.docx", "rb") as f:
        response.data = f.read()

    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = "attachment; filename=cv_pialroy.docx"
    return response  
  

# Run App
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
